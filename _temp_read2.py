import docx
import os

src_path = r"C:\Users\zero\.openclaw\media\inbound\视频主题_组长A_组员_BCDE_2---ee0768c5-84cc-4509-affa-66d732eac6f8.docx"
doc = docx.Document(src_path)

print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

print("\n=== 段落内容 ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text.strip()}")

print("\n=== 表格内容 ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip()[:30] for c in row.cells]
        print(f"行{ri}: {cells}")
